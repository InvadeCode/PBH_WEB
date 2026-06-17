from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page Numbers in Footer ---
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add "PurpleBlue House  |  Page X" 
run1 = footer_para.add_run('PurpleBlue House  |  Page ')
run1.font.size = Pt(8)
run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run1.font.name = 'Calibri'

# Insert auto page number field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run2 = footer_para.add_run()
run2._r.append(fldChar1)

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' PAGE '
run3 = footer_para.add_run()
run3._r.append(instrText)

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run4 = footer_para.add_run()
run4._r.append(fldChar2)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

PURPLE = RGBColor(0x5B, 0x2D, 0x8E)
DARK_PURPLE = RGBColor(0x2D, 0x1B, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = RGBColor(0xF5, 0xF0, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_purple_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_PURPLE
        run.font.name = 'Calibri'
    return h

def add_section_header(number, title):
    """Creates the purple numbered section header like the original"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(1.5)
    tbl.columns[1].width = Cm(15)
    
    # Number cell
    cell0 = tbl.cell(0, 0)
    set_cell_shading(cell0, '5B2D8E')
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(number)
    run0.font.color.rgb = WHITE
    run0.font.bold = True
    run0.font.size = Pt(12)
    run0.font.name = 'Calibri'
    
    # Title cell
    cell1 = tbl.cell(0, 1)
    set_cell_shading(cell1, 'E8DCF5')
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run('  ' + title)
    run1.font.color.rgb = DARK_PURPLE
    run1.font.bold = True
    run1.font.size = Pt(13)
    run1.font.name = 'Calibri'
    
    # Remove borders
    for cell in [cell0, cell1]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    
    doc.add_paragraph()

def add_styled_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_shading(cell, '5B2D8E')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Data rows
    for r_idx, row in enumerate(rows):
        bg = 'F9F9F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    
    doc.add_paragraph()

def add_note(text, prefix="For Blogs & Case Studies:"):
    p = doc.add_paragraph()
    run_prefix = p.add_run(prefix + ' ')
    run_prefix.font.bold = True
    run_prefix.font.size = Pt(10)
    run_prefix.font.color.rgb = PURPLE
    run_prefix.font.name = 'Calibri'
    run_text = p.add_run(text)
    run_text.font.italic = True
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = GREY
    run_text.font.name = 'Calibri'

def add_action_note(text):
    p = doc.add_paragraph()
    run_prefix = p.add_run('⚠ Action Required: ')
    run_prefix.font.bold = True
    run_prefix.font.size = Pt(10)
    run_prefix.font.color.rgb = RGBColor(0xD9, 0x73, 0x06)
    run_prefix.font.name = 'Calibri'
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.name = 'Calibri'

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PurpleBlue House')
run.font.size = Pt(28)
run.font.color.rgb = DARK_PURPLE
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SEO + AIO + GEO Content Template')
run.font.size = Pt(16)
run.font.color.rgb = PURPLE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Content Production Standard for Website Pages')
run.font.size = Pt(13)
run.font.color.rgb = GREY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Official Standard Operating Procedure\nJune 2026')
run.font.size = Pt(10)
run.font.color.rgb = GREY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('This document serves as a guiding framework for structuring and publishing web pages.\nIt is designed to support our teams in creating consistent, optimized content across blogs, case studies, and other pages.\nWe encourage all content creators to reference these best practices to ensure a unified standard.')
run.font.size = Pt(10)
run.font.color.rgb = GREY
run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# IMPLEMENTATION STATUS
# ============================================================
p = doc.add_paragraph()
run = p.add_run('Implementation Status')
run.font.size = Pt(20)
run.font.color.rgb = GREEN
run.font.bold = True
run.font.name = 'Calibri'
run2 = p.add_run('    LIVE')
run2.font.size = Pt(11)
run2.font.color.rgb = GREEN
run2.font.bold = True
run2.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('The advanced SEO / AIO / GEO framework has been ')
run.font.name = 'Calibri'
run2 = p.add_run('successfully implemented for all dynamic content (Case Studies and Blogs)')
run2.font.bold = True
run2.font.name = 'Calibri'
run3 = p.add_run(' on the PurpleBlue House website. The technical infrastructure is live and operational as of June 2026.')
run3.font.name = 'Calibri'

p = doc.add_paragraph('While all static pages carry foundational SEO, a highly advanced, custom-built engine now powers the Case Studies and Journal sections. This ensures that Google Search, AI-powered engines (ChatGPT, Perplexity, Google AI Overviews), and location-based search queries can all correctly discover, read and rank our most important thought leadership content.')

doc.add_paragraph()
h = doc.add_heading('Dynamic Content Engine — Case Studies & Journal Articles', level=2)
for run in h.runs:
    run.font.color.rgb = DARK_PURPLE
    run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('For ')
run.font.name = 'Calibri'
run2 = p.add_run('Case Studies')
run2.font.bold = True
run2.font.name = 'Calibri'
run3 = p.add_run(' and ')
run3.font.name = 'Calibri'
run4 = p.add_run('Journal Articles (Blogs)')
run4.font.bold = True
run4.font.name = 'Calibri'
run5 = p.add_run(', we have built a highly advanced, dynamic SEO engine directly into the Sanity CMS dashboard. The PBH content team now has direct control over the following fields for every single post:')
run5.font.name = 'Calibri'

add_styled_table(
    ['Field', 'Description', 'Validation'],
    [
        ['SEO Title', 'The title that appears in Google search results and browser tabs', 'Max 60 characters (warning enforced)'],
        ['Meta Description', 'The summary snippet shown below the title in search results', 'Max 155 characters (warning enforced)'],
        ['Primary Keyword', 'The single exact search phrase this page must rank for', 'Free text'],
        ['Page-Specific FAQs', 'Select 5–8 unique FAQs relevant to this specific article', 'Reference selector (linked to FAQ library)'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Automated Schema Markup: ')
run.font.bold = True
run.font.name = 'Calibri'
run2 = p.add_run('The website code automatically reads these Sanity inputs and injects proper JSON-LD structured data into every Case Study and Blog page:')
run2.font.name = 'Calibri'

doc.add_paragraph('• ItemPage schema — for Case Studies', style='List Bullet')
doc.add_paragraph('• Article schema — for Journal Articles', style='List Bullet')
doc.add_paragraph('• FAQPage schema — automatically generated when Page-Specific FAQs are added', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph('This structured data enables Google Rich Snippets (FAQ dropdowns in search results) and ensures AI engines can accurately cite PBH content when users search for queries like "best branding agency in India".')

add_action_note('The PBH content team must fill out the SEO Title, Meta Description, Primary Keyword and Page-Specific FAQs for every Case Study and Blog post before publishing. If these fields are left blank, the system will fall back to generic defaults, which reduces search performance.')

doc.add_page_break()

# ============================================================
# HOW TO USE THIS TEMPLATE
# ============================================================
add_purple_heading('How to Use This Template', level=1)

p = doc.add_paragraph('This document has two parts:')
doc.add_paragraph('Part 1 — The Template: Each section is a numbered field to fill out before writing a page.', style='List Bullet')
doc.add_paragraph('Part 2 — A Worked Example: A complete, filled-in example for a PBH client service page.', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Follow the order. Fill every field. Do not skip sections. The template exists because inconsistent content structure creates inconsistent search performance.')

add_styled_table(
    ['Page Type', 'When to Use'],
    [
        ['Blog', 'Thought leadership, how-to guides, industry explainers'],
        ['Case Study', 'Completed client projects with results'],
        ['Service Page', 'Core services offered by the client'],
        ['Product Page', 'Physical or digital products'],
        ['Location Page', 'City or region-specific service targeting'],
        ['Industry Page', 'Vertical-specific content (real estate, healthcare, etc.)'],
    ]
)

doc.add_page_break()

# ============================================================
# PART 1 — THE CONTENT TEMPLATE
# ============================================================
h = doc.add_heading('PART 1', level=1)
for run in h.runs:
    run.font.color.rgb = DARK_PURPLE
p = doc.add_paragraph()
run = p.add_run('The Content Template')
run.font.size = Pt(18)
run.font.color.rgb = DARK_PURPLE
run.font.name = 'Calibri'
doc.add_paragraph()

# 01
add_section_header('01', 'Page Objective')
doc.add_paragraph('What is this page supposed to achieve? Answer in 2–3 sentences. Be specific about the buyer action you want.')
add_styled_table(['Field', 'Instructions'], [
    ['Page Type', 'Select: Blog / Case Study / Service Page / Product Page / Location Page / Industry Page'],
    ['Objective Statement', 'Write what this page must do for the visitor and the business.\nExample: This page should help municipal corporations and smart city teams understand how to choose the right fabrication partner for bus shelter infrastructure.'],
])

# 02
add_section_header('02', 'Target Audience')
doc.add_paragraph('List the exact audience types this page speaks to. Be as specific as possible — avoid generic labels like \'businesses\' or \'clients\'.')
add_styled_table(['Field', 'Instructions'], [
    ['Primary Audience', 'Who is the single most important reader for this page?'],
    ['Secondary Audience(s)', 'List 3–6 secondary reader types, one per line.'],
])

# 03
add_section_header('03', 'Primary Keyword')
doc.add_paragraph('One keyword only. This is the single term the page must rank for. It must appear in the H1, SEO title, meta description, first paragraph and URL.')
add_styled_table(['Field', 'Instructions'], [
    ['Primary Keyword', 'Write one phrase only.\nExample: Brand identity agency in India'],
])
add_note('This field is directly editable in the Sanity CMS dashboard. The website automatically uses it in the JSON-LD structured data.')

# 04
add_section_header('04', 'Secondary Keywords')
doc.add_paragraph('5–8 supporting keywords. These should appear naturally in the body content — in headings, paragraphs and FAQs. Do not stuff them.')
add_styled_table(['Field', 'Instructions'], [
    ['Secondary Keywords', 'List 5–8 phrases, one per line.\nExample: Logo design agency India / Brand strategy firm / Visual identity designer / Brand naming agency'],
])

# 05
add_section_header('05', 'SEO Title')
doc.add_paragraph('Use the formula: Primary Keyword + Use Case or Buyer Intent + Geography or Context')
add_styled_table(['Field', 'Instructions'], [
    ['SEO Title', 'Write your title here. Keep it under 60 characters.\nExample: Brand Identity Agency in India for Startups and Growing Businesses'],
])
add_note('This field is directly editable in the Sanity CMS dashboard. The website automatically uses it as the browser tab title and the JSON-LD headline. Built-in validation warns if exceeding 60 characters.')

# 06
add_section_header('06', 'Meta Description')
doc.add_paragraph('One sentence. No keyword stuffing. Must be useful to someone scanning search results. Keep it under 155 characters.')
add_styled_table(['Field', 'Instructions'], [
    ['Meta Description', 'Write your meta description here.\nExample: Learn how PurpleBlue House builds brand identities for startups, SMEs and enterprise clients across India.'],
])
add_note('This field is directly editable in the Sanity CMS dashboard. Built-in validation warns if exceeding 155 characters.')

# 07
add_section_header('07', 'URL Slug')
doc.add_paragraph('Short, lowercase, hyphenated. Include the primary keyword. No dates in blog slugs unless the date is the point.')
add_styled_table(['Field', 'Instructions'], [
    ['URL Slug', '/section/keyword-phrase\nExample: /services/brand-identity-agency-india'],
])

# 08
add_section_header('08', 'H1 Heading')
doc.add_paragraph('One H1 per page. It must contain the primary keyword. It should be written for a human reader, not for a search engine.')
add_styled_table(['Field', 'Instructions'], [
    ['H1 Heading', 'Write your single H1 here.\nExample: Brand Identity Agency in India for Businesses Ready to Stand Out'],
])

# 09
add_section_header('09', 'Opening Answer Paragraph')
doc.add_paragraph('This is the most important paragraph on the page. It must directly answer the main topic in 80–120 words. It is the paragraph that AIO and GEO systems read first.')
doc.add_paragraph('Rules for the opening paragraph:')
doc.add_paragraph('Answer the topic directly in the first sentence.', style='List Bullet')
doc.add_paragraph('Include the primary keyword naturally.', style='List Bullet')
doc.add_paragraph('Mention the company name once.', style='List Bullet')
doc.add_paragraph('Do not start with a question.', style='List Bullet')
doc.add_paragraph('Do not start with a generic line like "In today\'s competitive world..."', style='List Bullet')

# 10
add_section_header('10', 'Suggested Page Structure')
doc.add_paragraph('Use the H2 structure below for every page. The structure should feel like a reference document, not a brochure.')
doc.add_paragraph('H2: What Is [Topic/Product/Service]?', style='List Bullet')
doc.add_paragraph('H2: Where Is It Used?', style='List Bullet')
doc.add_paragraph('H2: Why This Matters', style='List Bullet')
doc.add_paragraph('H2: Key Components / Features', style='List Bullet')
doc.add_paragraph('H2: Material and Technical Considerations', style='List Bullet')
doc.add_paragraph('H2: Process Followed by [Client Name]', style='List Bullet')
doc.add_paragraph('H2: Comparison Table', style='List Bullet')
doc.add_paragraph('H2: What Buyers Should Check Before Finalising a Vendor', style='List Bullet')
doc.add_paragraph('H2: Why [Client Name]', style='List Bullet')
doc.add_paragraph('H2: FAQs', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Not every H2 will apply to every page. Use your judgment. If a section adds no useful information for the target reader, remove it.')

# 11
add_section_header('11', 'Mandatory FAQ Section')
doc.add_paragraph('Every page must have a FAQ section. Minimum 5 FAQs. Ideal: 6–8 FAQs. Write real questions that buyers would actually ask.')
add_styled_table(
    ['FAQ #', 'Required Question Theme', 'Target Length'],
    [
        ['FAQ 1', 'What is [main topic/product]?', '40–70 words'],
        ['FAQ 2', 'Where is [product/service] commonly used?', '40–70 words'],
        ['FAQ 3', 'Which material or approach is best?', '40–70 words'],
        ['FAQ 4', 'How long does execution take?', '30–50 words'],
        ['FAQ 5', 'What should buyers check before choosing a vendor?', '40–60 words'],
        ['FAQ 6', 'Does [client] provide customisation?', '30–50 words'],
        ['FAQ 7', 'Does [client] handle installation or full delivery?', '30–50 words'],
        ['FAQ 8', 'Does [client] work across India or internationally?', '30–50 words'],
    ]
)
add_note('FAQs can be selected per-page directly from the Sanity CMS. When added, the website automatically generates FAQPage JSON-LD schema, enabling Google Rich Snippet display.')

# 12
add_section_header('12', 'Internal Links')
doc.add_paragraph('Every page must link to 3–6 related pages. Internal links help readers navigate and help search engines understand site structure.')
add_styled_table(['Field', 'Instructions'], [
    ['Internal Links', 'List 3–6 page names and their URLs.\nExample: /services/logo-design | /services/brand-strategy | /case-studies | /contact'],
])

# 13
add_section_header('13', 'Image Requirements')
doc.add_paragraph('Each page should have 2–5 images where possible. Every image must have a descriptive file name and alt text.')
add_styled_table(
    ['Element', 'Format', 'Example'],
    [
        ['File name', 'lowercase-with-hyphens.jpg', 'brand-identity-design-purple-blue-house.jpg'],
        ['Alt text', 'Describe the image with keyword', 'Brand identity system designed by PurpleBlue House for a Mumbai-based startup'],
    ]
)

# 14
add_section_header('14', 'Schema Markup')
doc.add_paragraph('Required structured data per page type:')
add_styled_table(
    ['Page Type', 'Required Schemas'],
    [
        ['Service Page', 'Service, FAQ, Breadcrumb, Organization'],
        ['Blog', 'Article, FAQ, Breadcrumb'],
        ['Case Study', 'ItemPage, FAQ, Breadcrumb'],
    ]
)
add_note('Article, ItemPage and FAQPage schemas are automatically injected by the website code based on CMS inputs. No manual implementation required.')

# 15
add_section_header('15', 'Writing Rules')
doc.add_paragraph('No filler words, no marketing fluff.', style='List Bullet')
doc.add_paragraph('Paragraphs: 3–4 sentences maximum.', style='List Bullet')
doc.add_paragraph('Every sentence must state a fact, explain a process or answer a question.', style='List Bullet')
doc.add_paragraph('Avoid: "In today\'s world…", "We are committed to…", "Best-in-class…"', style='List Bullet')
doc.add_paragraph('Write as if explaining to a knowledgeable buyer, not selling to them.', style='List Bullet')

# 16
add_section_header('16', 'Publishing Checklist')
add_styled_table(
    ['#', 'Check', 'Done?'],
    [
        ['1', 'H1 contains primary keyword', '☐'],
        ['2', 'SEO Title under 60 characters', '☐'],
        ['3', 'Meta Description under 155 characters', '☐'],
        ['4', 'Primary keyword in first 100 words', '☐'],
        ['5', 'URL slug contains keyword', '☐'],
        ['6', '5–8 FAQs added', '☐'],
        ['7', '3–6 internal links included', '☐'],
        ['8', 'All images have descriptive alt text', '☐'],
        ['9', 'No filler language', '☐'],
        ['10', 'Opening paragraph answers the topic directly', '☐'],
    ]
)

# 17
add_section_header('17', 'CTA Format')
doc.add_paragraph('Every page must end with a clear call to action. Use this format:')
p = doc.add_paragraph()
run = p.add_run('Action verb + outcome + contact method.')
run.font.bold = True
run.font.name = 'Calibri'
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Example: ')
run.font.italic = True
run.font.name = 'Calibri'
run2 = p.add_run('Building or redesigning your brand? Talk to PurpleBlue House about brand strategy, identity design and execution. Contact us to discuss your project.')
run2.font.italic = True
run2.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# QUICK REFERENCE
# ============================================================
add_purple_heading('Quick Reference: Content Standards', level=1)

h = doc.add_heading('What to Never Write', level=2)
for run in h.runs:
    run.font.color.rgb = DARK_PURPLE
add_styled_table(
    ['❌ Avoid', '✅ Write Instead'],
    [
        ['We are a leading agency', 'We work with startups and enterprise brands across India'],
        ['We provide quality solutions', 'Our process covers strategy, design and delivery'],
        ['We are committed to excellence', 'State what you actually do'],
        ['Best-in-class services', 'Describe the actual service'],
        ['We offer end-to-end solutions', 'List what the end-to-end scope actually includes'],
        ['We have a passionate team', "Mention the team's experience or specific capabilities"],
    ]
)

h = doc.add_heading('Keyword Placement Rules', level=2)
for run in h.runs:
    run.font.color.rgb = DARK_PURPLE
add_styled_table(
    ['Location', 'Requirement'],
    [
        ['H1 Heading', 'Primary keyword must appear'],
        ['SEO Title', 'Primary keyword must appear, preferably near the start'],
        ['Meta Description', 'Primary keyword must appear once'],
        ['Opening Paragraph', 'Primary keyword must appear in first 100 words'],
        ['URL Slug', 'Primary keyword must appear'],
        ['At least one H2', 'Primary keyword should appear naturally'],
        ['FAQs', 'Primary and secondary keywords should appear naturally'],
        ['Image Alt Text', 'Relevant keyword included where it fits naturally'],
    ]
)

h = doc.add_heading('Page Length Guidelines', level=2)
for run in h.runs:
    run.font.color.rgb = DARK_PURPLE
add_styled_table(
    ['Page Type', 'Minimum Word Count', 'Ideal Word Count'],
    [
        ['Blog', '800 words', '1,200–2,000 words'],
        ['Case Study', '600 words', '900–1,500 words'],
        ['Service Page', '700 words', '1,000–1,800 words'],
        ['Product Page', '400 words', '600–1,200 words'],
        ['Location Page', '600 words', '800–1,400 words'],
        ['Industry Page', '700 words', '1,000–1,800 words'],
    ]
)

# Save
output_path = '/Users/shravanikhurpe/Desktop/PBH_WEB/PBH_SEO_AIO_GEO_Updated_v2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
